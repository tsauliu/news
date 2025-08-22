#%%
# Step 3: Generate English Word document from three markdown files
# This script translates three markdown files and combines them into one English Word document

from docx import Document
from parameters import friday_date
from models import deepseek_model,gemini_model
import pandas as pd
import re
import os
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
from tqdm import tqdm

# Input and output directories
input_dir = f'data/6_final_mds'
output_dir = f'data/6_final_mds'
docx_dir = 'data/7_docx'

# Input markdown files
sellside_md = f'{input_dir}/{friday_date}_sellside_highlights.md'
takeaway_md = f'{input_dir}/{friday_date}_key_takeaway.md'
detailed_md = f'{input_dir}/{friday_date}_detailed_news.md'

# Output English markdown files
sellside_eng_md = f'{output_dir}/{friday_date}_sellside_highlights_english.md'
takeaway_eng_md = f'{output_dir}/{friday_date}_key_takeaway_english.md'
detailed_eng_md = f'{output_dir}/{friday_date}_detailed_news_english.md'

os.makedirs(output_dir, exist_ok=True)
os.makedirs(docx_dir, exist_ok=True)

# Translation prompt for DeepSeek (line by line translation)
translation_prompt = """You are a professional translator specializing in automotive industry content. 
Translate the following Chinese text to natural, accurate English. 
Context: This is automotive industry news, research reports, and market analysis.
Requirements:
1. Maintain professional automotive terminology
2. Keep the original meaning and tone
3. Use natural, native English expressions
4. Keep company names, brand names, or proper nouns accurate
5. Keep technical terms accurate
6. Only return the translated text, no explanations or additional content
7. IMPORTANT: Your output must be ENGLISH ONLY - no Chinese characters allowed in the response
8. If the input is already in English, return it unchanged
9. Never include any Chinese text in your response under any circumstances

Specific translation guidelines for key terms:
- 商业落地 → Commercialization
- 核心技术 → New Technology
- 政策监管 → Policy Regulation
- 企业战略 → Corporate Strategy
- 硬件设备 → Hardware
- 资本动向 → Capital Trends

Please translate the following text:
"""

# Special translation prompt for Gemini (full file translation)
gemini_translation_prompt = """You are a professional translator specializing in automotive industry content. 
Translate the following Chinese text to natural, accurate English. 
Context: This is automotive industry news, research reports, and market analysis.

CRITICAL FORMATTING REQUIREMENTS:
1. ONLY translate text content - DO NOT modify any formatting or structure
2. Keep ALL HTML comments and tags EXACTLY as they are (e.g., <!-- WORD_STYLE: bullet -->)
3. Keep ALL links COMPLETELY unchanged (URLs, link text, markdown link syntax)
4. Preserve ALL markdown formatting (headers, bullet points, numbering, etc.)
5. Maintain the exact line breaks and spacing structure
6. Do not add, remove, or modify any special characters or symbols

Translation Requirements:
1. Maintain professional automotive terminology
2. Keep the original meaning and tone
3. Use natural, native English expressions
4. Keep company names, brand names, or proper nouns accurate
5. Keep technical terms accurate
6. IMPORTANT: Your output must be ENGLISH ONLY - no Chinese characters allowed in the response
7. If text is already in English, keep it unchanged
8. Never include any Chinese text in your response under any circumstances

Specific translation guidelines for key terms:
- 商业落地 → Commercialization
- 核心技术 → New Technology
- 政策监管 → Policy Regulation
- 企业战略 → Corporate Strategy
- 硬件设备 → Hardware
- 资本动向 → Capital Trends

REMEMBER: Only translate the text content, preserve all formatting, tags, and links exactly as they are.

Please translate the following text:
"""

# Thread lock for print statements
print_lock = threading.Lock()

def safe_print(message):
    """Thread-safe print function"""
    with print_lock:
        print(message)

def translate_with_gemini(text):
    """Translate entire file content using Gemini model with retries"""
    if not text or len(text.strip()) < 3:
        return text
    
    import time
    max_retries = 3
    
    for attempt in range(max_retries):
        try:
            translated = gemini_model(gemini_translation_prompt, text)
            return translated.strip()
        except Exception as e:
            if attempt < max_retries - 1:
                wait_time = 5 * (2 ** attempt)  # 5, 10, 20 seconds
                safe_print(f"Gemini translation error (attempt {attempt + 1}/{max_retries}): {e}")
                safe_print(f"Retrying in {wait_time} seconds...")
                time.sleep(wait_time)
            else:
                safe_print(f"Gemini translation failed after {max_retries} attempts: {e}")
                return ""  # Return blank instead of Chinese text

def translate_with_deepseek(text):
    """Translate text using DeepSeek model with retries"""
    if not text or len(text.strip()) < 3:
        return text
    
    import time
    max_retries = 3
    
    for attempt in range(max_retries):
        try:
            translated = deepseek_model(translation_prompt, text)
            return translated.strip()
        except Exception as e:
            if attempt < max_retries - 1:
                wait_time = 5 * (2 ** attempt)  # 5, 10, 20 seconds
                safe_print(f"DeepSeek translation error (attempt {attempt + 1}/{max_retries}): {e}")
                safe_print(f"Retrying in {wait_time} seconds...")
                time.sleep(wait_time)
            else:
                safe_print(f"DeepSeek translation failed after {max_retries} attempts: {e}")
                return ""  # Return blank instead of Chinese text

def translate_file_with_gemini(input_file, output_file):
    """Translate entire file using Gemini"""
    print(f"Translating {input_file} with Gemini...")
    
    if not os.path.exists(input_file):
        print(f"Warning: File {input_file} not found, skipping...")
        return
    
    # Read the entire file
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Translate the entire content
    translated_content = translate_with_gemini(content)
    
    # Write to output file
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(translated_content)
    
    print(f"Gemini translation completed: {output_file}")

def parse_markdown_to_dataframe(content):
    """Parse markdown content and return a dataframe with index, style, and text"""
    lines = content.split('\n')
    data = []
    index = 0
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines but keep track of them for proper spacing
        if not line:
            data.append({'index': index, 'style': 'empty', 'text': ''})
            index += 1
            i += 1
            continue
        
        # Check for Word style annotations
        if line.startswith('<!-- WORD_STYLE:'):
            # Extract the style
            style_match = re.search(r'<!-- WORD_STYLE: (.+?) -->', line)
            if style_match:
                style = style_match.group(1).strip()
                i += 1  # Move to next line for content
                
                # Get the content line(s)
                if i < len(lines):
                    content_line = lines[i].strip()
                    data.append({'index': index, 'style': style, 'text': content_line})
                    index += 1
                else:
                    # Style annotation without content
                    data.append({'index': index, 'style': style, 'text': ''})
                    index += 1
            i += 1
        else:
            # Line without style annotation - treat as normal paragraph
            data.append({'index': index, 'style': 'normal_paragraph', 'text': line})
            index += 1
            i += 1
    
    return pd.DataFrame(data)

def translate_row(row):
    """Translate a single row from the dataframe"""
    index, style, text = row['index'], row['style'], row['text']
    
    # Determine if we should translate this content
    should_translate = style not in ['link', 'page_break']
    
    if should_translate and text and len(text.strip()) >= 3:
        safe_print(f"Translating row {index}: {text[:50]}...")
        translated_text = translate_with_deepseek(text)
        return {'index': index, 'style': style, 'text': text, 'translated_text': translated_text}
    else:
        return {'index': index, 'style': style, 'text': text, 'translated_text': text}

def translate_detailed_news_with_deepseek(input_file, output_file):
    """Translate detailed news file line by line using DeepSeek"""
    print(f"Translating {input_file} with DeepSeek (line by line)...")
    
    if not os.path.exists(input_file):
        print(f"Warning: File {input_file} not found, skipping...")
        return
    
    # Read and parse the file
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    df = parse_markdown_to_dataframe(content)
    print(f"Parsed {len(df)} items from markdown file")
    
    # Filter rows that need translation
    rows_to_translate = df[
        (df['style'] != 'link') & 
        (df['style'] != 'page_break') & 
        (df['style'] != 'empty') & 
        (df['text'].str.len() >= 3)
    ]
    
    print(f"Found {len(rows_to_translate)} items to translate")
    
    # Use ThreadPoolExecutor for multi-threading
    max_workers = 50  # Adjust based on your system and API limits
    translated_results = []
    
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        # Submit all translation tasks
        future_to_row = {
            executor.submit(translate_row, row): idx 
            for idx, row in rows_to_translate.iterrows()
        }
        
        # Process completed translations with progress bar
        with tqdm(total=len(future_to_row), desc="Translating with DeepSeek") as pbar:
            for future in as_completed(future_to_row):
                try:
                    result = future.result()
                    translated_results.append(result)
                    pbar.update(1)
                except Exception as e:
                    safe_print(f"Translation task failed: {e}")
                    pbar.update(1)
    
    # Create translated dataframe
    translated_df = pd.DataFrame(translated_results)
    
    # Merge with original dataframe
    df = df.merge(
        translated_df[['index', 'translated_text']], 
        on='index', 
        how='left'
    )
    
    # Fill missing translations with blank text instead of Chinese
    df['translated_text'] = df['translated_text'].fillna("")
    
    print(f"Translation completed. Processed {len(translated_results)} items")
    
    # Reconstruct the markdown file
    output_lines = []
    for _, row in df.iterrows():
        if row['style'] == 'empty':
            output_lines.append('')
        else:
            output_lines.append(f'<!-- WORD_STYLE: {row["style"]} -->')
            output_lines.append(row['translated_text'])
    
    # Write to output file
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write('\n'.join(output_lines))
    
    print(f"DeepSeek translation completed: {output_file}")

def process_markdown_file(file_path, doc):
    """Process a markdown file and add content to the Word document"""
    print(f"Processing: {file_path}")
    
    if not os.path.exists(file_path):
        print(f"Warning: File {file_path} not found, skipping...")
        return
    
    # Read the markdown file
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Split content by lines and process
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Check for Word style annotations
        if line.startswith('<!-- WORD_STYLE:'):
            # Extract the style
            style_match = re.search(r'<!-- WORD_STYLE: (.+?) -->', line)
            if style_match:
                style = style_match.group(1).strip()
                i += 1  # Move to next line for content
                
                # Get the content line(s)
                if i < len(lines):
                    content_line = lines[i].strip()
                    
                    # Apply formatting based on style
                    if style == 'heading_level_1':
                        # Remove markdown # symbols and add as Word heading
                        heading_text = content_line.replace('#', '').strip()
                        doc.add_heading(heading_text, level=1)
                    
                    elif style == 'heading_level_2':
                        # Remove markdown ## symbols and add as Word heading
                        heading_text = content_line.replace('#', '').strip()
                        doc.add_heading(heading_text, level=2)
                    
                    elif style == 'heading_level_3':
                        # Remove markdown ### symbols and add as Word heading
                        heading_text = content_line.replace('#', '').strip()
                        doc.add_heading(heading_text, level=3)
                    
                    elif style == 'summarytitle':
                        doc.add_paragraph('')
                        doc.add_paragraph(content_line, style='summarytitle')
                    
                    elif style == 'bullet':
                        doc.add_paragraph(content_line, style='bullet')
                    
                    elif style == 'link':
                        doc.add_paragraph(content_line, style='link')
                    
                    elif style == 'author':
                        doc.add_paragraph(content_line, style='author')
                    
                    elif style == 'normal_paragraph':
                        if content_line:  # Only add non-empty paragraphs
                            doc.add_paragraph(content_line)
                            doc.add_paragraph('')  # Add empty paragraph after content
                    
                    elif style == 'page_break':
                        doc.add_page_break()
                        i += 1  # Skip the content line for page breaks
                        continue
                    
                    else:
                        # Unknown style, treat as normal paragraph
                        if content_line:  # Only add non-empty paragraphs
                            doc.add_paragraph(content_line)
            
            i += 1
        else:
            # Line without style annotation - treat as normal paragraph if not empty
            if line:
                doc.add_paragraph(line)
            i += 1

def main():
    print("=== Step 1: Translating Sellside Highlights with Gemini ===")
    translate_file_with_gemini(sellside_md, sellside_eng_md)
    
    print("\n=== Step 2: Translating Key Takeaway with Gemini ===")
    translate_file_with_gemini(takeaway_md, takeaway_eng_md)
    
    print("\n=== Step 3: Translating Detailed News with DeepSeek ===")
    translate_detailed_news_with_deepseek(detailed_md, detailed_eng_md)
    
    print("\n=== Step 4: Assembling English Word Document ===")
    # Create new Word document using template
    doc = Document('news_template.docx')
    doc.add_page_break()
    
    # Process the three English markdown files in order
    print("Processing English Sellside Highlights...")
    process_markdown_file(sellside_eng_md, doc)
    
    # Add page break after sellside highlights
    doc.add_page_break()
    
    print("Processing English Key News Takeaway...")
    process_markdown_file(takeaway_eng_md, doc)
    
    # Add page break after key takeaway
    doc.add_page_break()
    
    # # Add Table of Contents
    # print("Adding Table of Contents...")
    # doc.add_heading('Table of Contents', level=1)
    # doc.add_page_break()
    
    print("Processing English Detailed News...")
    process_markdown_file(detailed_eng_md, doc)
    
    # Save the document
    output_file = f'{docx_dir}/Autonomous Driving AI News Summary {friday_date.replace("-", " ")}_ENG.docx'
    doc.save(output_file)
    
    print(f"\nEnglish Word document generated: {output_file}")
    print("Document creation completed!")
    
    print("\nGenerated files:")
    print(f"1. {sellside_eng_md}")
    print(f"2. {takeaway_eng_md}")
    print(f"3. {detailed_eng_md}")
    print(f"4. {output_file}")

if __name__ == "__main__":
    main() 