# Step 2: Generate Chinese Word document from three reviewed markdown files
# This script reads three markdown files and combines them into one Word document with proper formatting

from docx import Document
from parameters import friday_date
import re
import os

# Input markdown files
input_dir = f'data/6_final_mds'
sellside_md = f'{input_dir}/{friday_date}_sellside_highlights.md'
takeaway_md = f'{input_dir}/{friday_date}_key_takeaway.md'
detailed_md = f'{input_dir}/{friday_date}_detailed_news.md'

# Create new Word document using template
doc = Document('news_template.docx')

def process_markdown_file(file_path, doc):
    """Process a markdown file and add content to the Word document"""
    print(f"Processing: {file_path}")
    
    if not os.path.exists(file_path):
        print(f"Warning: File {file_path} not found, skipping...")
        return
    
    # Read the markdown file
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Split content by lines and process
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Check for Word style annotations
        if line.startswith('<!-- WORD_STYLE:'):
            # Extract the style
            style_match = re.search(r'<!-- WORD_STYLE: (.+?) -->', line)
            if style_match:
                style = style_match.group(1).strip()
                i += 1  # Move to next line for content
                
                # Get the content line(s)
                if i < len(lines):
                    content_line = lines[i].strip()
                    
                    # Apply formatting based on style
                    if style == 'heading_level_1':
                        # Remove markdown # symbols and add as Word heading
                        heading_text = content_line.replace('#', '').strip()
                        doc.add_heading(heading_text, level=1)
                    
                    elif style == 'heading_level_2':
                        # Remove markdown ## symbols and add as Word heading
                        heading_text = content_line.replace('#', '').strip()
                        doc.add_heading(heading_text, level=2)
                    
                    elif style == 'heading_level_3':
                        # Remove markdown ### symbols and add as Word heading
                        heading_text = content_line.replace('#', '').strip()
                        doc.add_heading(heading_text, level=3)
                    
                    elif style == 'summarytitle':
                        doc.add_paragraph('')
                        doc.add_paragraph(content_line, style='summarytitle')
                    
                    elif style == 'bullet':
                        doc.add_paragraph(content_line, style='bullet')
                    
                    elif style == 'link':
                        doc.add_paragraph(content_line, style='link')
                    
                    elif style == 'author':
                        doc.add_paragraph(content_line, style='author')
                    
                    elif style == 'normal_paragraph':
                        doc.add_paragraph(content_line)
                        doc.add_paragraph('')  # Add empty paragraph after content
                    
                    elif style == 'page_break':
                        doc.add_page_break()
                        i += 1  # Skip the content line for page breaks
                        continue
                    
                    else:
                        # Unknown style, treat as normal paragraph
                        doc.add_paragraph(content_line)
            
            i += 1
        else:
            # Line without style annotation - treat as normal paragraph if not empty
            if line:
                doc.add_paragraph(line)
            i += 1

# Process the three markdown files in order
print("Step 1: Processing Sellside Highlights...")
process_markdown_file(sellside_md, doc)

# Add page break after sellside highlights
doc.add_page_break()

print("Step 2: Processing Key News Takeaway...")
process_markdown_file(takeaway_md, doc)

# Add page break after key takeaway
doc.add_page_break()

# Add Table of Contents
print("Step 3: Adding Table of Contents...")
doc.add_heading('Table of Contents', level=1)
doc.add_page_break()

print("Step 4: Processing Detailed News...")
process_markdown_file(detailed_md, doc)

# Save the document
output_dir = 'data/7_docx'
os.makedirs(output_dir, exist_ok=True)
output_file = f'{output_dir}/{friday_date}_weekly_news.docx'
doc.save(output_file)

print(f"Chinese Word document generated: {output_file}")
print("Document creation completed!") 