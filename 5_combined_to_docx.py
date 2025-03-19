# combine the summary and detailed news into a docx file
# %%
from docx import Document
import pandas as pd
import datetime
from parameters import friday_date

doc = Document('news_template.docx')
summary_md=open(f'data/3_summary_mds/{friday_date}_summary.md', 'r', encoding='utf-8').read()
## add the key takeaway for the week
doc.add_heading('Key takeaway for Week – Mar 19', level=1)
doc.add_paragraph('')
# Parse the summary markdown and add headings and paragraphs
lines = summary_md.strip().split('\n')
for line in lines:
    if line.startswith('## '):
        doc.add_paragraph('')
        doc.add_paragraph(line[3:], style='summarytitle')
    elif line.startswith('# '):
        # Add level 3 heading (title3)
        doc.add_paragraph(line[2:], style='bullet')

doc.add_page_break()
doc.add_heading('Detailed News for Week – Mar 19', level=1)
doc.add_paragraph('')
combined_md=open('data/2_combined_mds/2025-03-19_merged_news.md', 'r', encoding='utf-8').read()
lines = combined_md.strip().split('\n')
# Convert lines to a dataframe

news_data = []

for line in lines:
    if line.startswith('title: '):
        news_data.append({'title': line[7:]})
    elif line.startswith('link: ') and news_data:
        news_data[-1]['link'] = line[6:]
    elif line.startswith('sector: ') and news_data:
        news_data[-1]['sector'] = line[8:].split('、')[0]
    elif line.startswith('author: ') and news_data:
        news_data[-1]['author'] = line[8:]
    elif line.startswith('date: ') and news_data:
        news_data[-1]['date'] = line[6:]
    elif line.startswith('content: ') and news_data:
        news_data[-1]['content'] = line[9:]

# Create dataframe
news_df = pd.DataFrame(news_data)
c1=news_df.sector!='其他'
news_df=news_df[c1].sort_values(by=['sector','date'], ascending=False)

# Loop through the dataframe to write to doc
for sector in ['核心技术','商业落地','政策监管','企业战略','硬件设备','数据与地图','资本动向']:
    news_df_sector=news_df[news_df.sector==sector]
    if news_df_sector.empty:
        continue
    doc.add_heading(sector, level=2)
    for _, row in news_df_sector.iterrows():
        doc.add_paragraph('')
        doc.add_heading(row['title'], level=3)
    
        if 'link' in row and not pd.isna(row['link']):
            doc.add_paragraph(row['link'], style='link')
        
        if 'author' in row and not pd.isna(row['author']):
            if 'date' in row and not pd.isna(row['date']):
                doc.add_paragraph(row['date'] + ' ' + row['author'], style='author')
            else:
                doc.add_paragraph(row['author'], style='author')
        
        if 'content' in row and not pd.isna(row['content']):
            doc.add_paragraph(row['content'])
            doc.add_paragraph('')

doc.save(f'data/{friday_date}_weekly_news.docx')
