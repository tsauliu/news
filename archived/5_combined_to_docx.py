# combine the summary and detailed news into a docx file
# %%
from pdfreport.run import auto_weekly_reports
auto_weekly_reports()

from docx import Document
import pandas as pd
import datetime,os
from parameters import friday_date,sector_list
import shutil

doc = Document('news_template.docx')

# research reports for the week
doc.add_heading(f'Sellside highlights for Week – {friday_date}', level=1)

raw_path=f'./pdfreport/01 raw/{friday_date}'
cdn_path=f'./pdfreport/cdn/{friday_date}'

os.makedirs(cdn_path, exist_ok=True)

# Sort files by date in filename (yyyy-mm-dd format)
for file in sorted(os.listdir(raw_path), key=lambda x: x.split('-')[0:3] if '-' in x else x, reverse=True):
    if file.endswith('.pdf'):
        print(file)
        ds_summary=open(f'./pdfreport/04 summary/{friday_date}_ds/{file.replace('.pdf', '.md')}', 'r', encoding='utf-8').read()
        # gemini_summary=open(f'./pdfreport/04 summary/{friday_date}_gemini/{file.replace('.pdf', '.md')}', 'r', encoding='utf-8').read()
        for summary in [ds_summary]:
            lines = summary.strip().split('\n')
            for line in lines:
                if line.startswith('**'):
                    doc.add_paragraph('')
                    doc.add_paragraph(line.replace('**','').strip(), style='summarytitle')
                elif len(line) > 10:
                    doc.add_paragraph(line.replace('*','').replace('**','').replace('- ','').replace('#','').strip(), style='bullet')
        
        parts = file.split('-')
        if len(parts) > 1:
            file_id = parts[-1].replace('.pdf', '')
            new_filename = f"{file_id}.pdf"
            source_path = os.path.join(raw_path, file)
            destination_path = os.path.join(cdn_path, new_filename)
            shutil.copy2(source_path, destination_path)
            # doc.add_paragraph(f'', style='link')
            doc.add_paragraph(f'https://auto.bda-news.com/{friday_date}/{file_id}.pdf', style='link')
            

# %% add the key takeaway for the week

doc.add_page_break()
summary_md=open(f'data/5_summary_mds/{friday_date}_summary.md', 'r', encoding='utf-8').read()

doc.add_heading(f'Key News takeaway for Week – {friday_date}', level=1)
# doc.add_paragraph('')
# Parse the summary markdown and add headings and paragraphs
lines = summary_md.strip().split('\n')
for line in lines:
    if line.startswith('##'):
        doc.add_paragraph('')
        doc.add_paragraph(line[2:].replace(' ',''), style='summarytitle')
    elif len(line) > 10 and line.startswith('#'):
        doc.add_paragraph(line.replace('*','').replace('**','').replace('- ','').replace('#','').replace(' ',''), style='bullet')

doc.add_page_break()

#%% add the detailed news for the week
doc.add_heading(f'Table of Contents', level=1)
doc.add_page_break()
doc.add_heading(f'Detailed News for Week – {friday_date}', level=1)
doc.add_paragraph('')
combined_md=open(f'data/4_combined_mds/{friday_date}_combined_news.md', 'r', encoding='utf-8').read()
lines = combined_md.strip().split('\n')

# Convert lines to a dataframe
news_data = []

for line in lines:
    if line.startswith('title: '):
        news_data.append({'title': line[7:]})
    elif line.startswith('link: ') and news_data:
        news_data[-1]['link'] = line[6:]
    elif line.startswith('sector: ') and news_data:
        news_data[-1]['sector'] = line[8:].split('、')[0]
    elif line.startswith('author: ') and news_data:
        news_data[-1]['author'] = line[8:]
    elif line.startswith('date: ') and news_data:
        news_data[-1]['date'] = line[6:]
    elif line.startswith('content: ') and news_data:
        news_data[-1]['content'] = line[9:]

# Create dataframe
news_df = pd.DataFrame(news_data)
c1=news_df.sector!='其他'
news_df=news_df[c1].sort_values(by=['sector','date'], ascending=False)

# Loop through the dataframe to write to doc
for sector in sector_list:
    news_df_sector=news_df[news_df.sector==sector]
    if news_df_sector.empty:
        continue
    doc.add_heading(sector, level=2)
    for _, row in news_df_sector.iterrows():
        doc.add_paragraph('')
        doc.add_heading(row['title'], level=3)
    
        if 'link' in row and not pd.isna(row['link']):
            doc.add_paragraph(row['link'], style='link')
        
        if 'author' in row and not pd.isna(row['author']):
            if 'date' in row and not pd.isna(row['date']):
                doc.add_paragraph(row['date'] + ' ' + row['author'], style='author')
            else:
                doc.add_paragraph(row['author'], style='author')
        
        if 'content' in row and not pd.isna(row['content']):
            doc.add_paragraph(row['content'])
            doc.add_paragraph('')

doc.save(f'data/{friday_date}_weekly_news.docx')
